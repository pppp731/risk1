"""
按PDF 10条要求，基于版本1111内容，生成完整技术报告。
保持原有样式体系（Heading 1 / Heading 2 / Normal / Caption）
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

# 复用版本1111的样式体系（读取模板）
doc = Document(r'C:\Users\23994\risk_sentinel\版本1111.docx')

# 清空所有内容
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)
for t in doc.tables:
    t._element.getparent().remove(t._element)
for s in doc.sections:
    for p in s.header.paragraphs:
        p.clear()
    for p in s.footer.paragraphs:
        p.clear()

# ===== 页面设置（保留版本1111的边距）=====
s = doc.sections[0]
s.top_margin = Emu(1332230)
s.bottom_margin = Emu(1259840)
s.left_margin = Emu(1007745)
s.right_margin = Emu(935990)

# ===== 计数器 =====
class Ctr:
    def __init__(self): self.ch = 1; self.t = 0; self.f = 0; self.eq = 0
    def chapter(self, n): self.ch = n; self.t = 0; self.f = 0; self.eq = 0
    def next_t(self): self.t += 1; return f"{self.ch}-{self.t}"
    def next_f(self): self.f += 1; return f"{self.ch}-{self.f}"
    def next_e(self): self.eq += 1; return f"({self.ch}-{self.eq})"
ctr = Ctr()

# ===== 辅助函数（使用版本1111的样式名）=====
def add_p(text, style='Normal', align=None, sb=0, sa=0, ls=1.5, indent=None):
    p = doc.add_paragraph(text, style=style)
    if align is not None: p.alignment = align
    pf = p.paragraph_format
    pf.space_before=Pt(sb); pf.space_after=Pt(sa); pf.line_spacing=ls
    if indent: pf.first_line_indent = Cm(indent)
    return p

def body(text): return add_p(text, 'Normal', None, 1, 1, 1.5, 0.74)
def body_ni(text): return add_p(text, 'Normal', None, 1, 1, 1.5, 0)
def h1(text, ch=None):
    if ch: ctr.chapter(ch)
    return add_p(text, 'Heading 1', None, 18, 10, 1.3, 0)
def h2(text): return add_p(text, 'Heading 2', None, 12, 6, 1.3, 0)
def h3(text): return add_p(text, 'Normal', None, 8, 4, 1.3, 0)  # 正文加粗做三级标题
def page_break(): doc.add_page_break()

def add_table(headers, rows, caption_text, cw=None):
    num = ctr.next_t()
    add_p(f'表{num}  {caption_text}', 'Caption', WD_ALIGN_PARAGRAPH.CENTER, 8, 2, 1.1, 0)
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>'))
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    if cw:
        for i,w in enumerate(cw):
            for row in t.rows: row.cells[i].width=Cm(w)
    return t

def add_figure(img_path, caption_text, width_inches=5.0):
    if os.path.exists(img_path):
        num = ctr.next_f()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format; pf.space_before=Pt(8); pf.space_after=Pt(2)
        p.add_run().add_picture(img_path, width=Inches(width_inches))
        add_p(f'图{num}  {caption_text}', 'Caption', WD_ALIGN_PARAGRAPH.CENTER, 0, 10, 1.1, 0)

def add_equation(eq_text, eq_num=None):
    if eq_num is None: eq_num = ctr.next_e()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before=Pt(6); pf.space_after=Pt(6); pf.line_spacing=1.3
    p.add_run(eq_text)
    p.add_run(f'    {eq_num}')
    return p

# ===== 页眉页脚 =====
for sec in doc.sections:
    hdr = sec.header; hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.add_run('智能风险监控Agent系统——技术报告')
    ftr = sec.footer; ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run('— ')
    fp._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    fp._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    fp._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    fp.add_run(' —')

# ==================== 封面 ====================
for _ in range(6): body_ni('')
add_p('智能风险监控Agent系统', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, 30, 0, 2.0, 0)
add_p('——技术报告', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, 10, 30, 1.5, 0)
for _ in range(2): body_ni('')
add_p('基于编排式Agent架构的企业风险管理平台', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, 5, 5, 1.5, 0)
add_p('以IF椰子水品牌为应用案例', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, 5, 5, 1.5, 0)
for _ in range(2): body_ni('')
add_p(f'技术栈：Python / Flask / NLP / Chart.js / yfinance', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, 3, 3, 1.5, 0)
add_p('代码规模：21个模块 · 12,000+行代码 · 16个API端点', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, 3, 3, 1.5, 0)
for _ in range(3): body_ni('')
add_p('作者：何培君', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, 3, 3, 1.5, 0)
add_p('学号：2024044048', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, 3, 3, 1.5, 0)
add_p(f'提交日期：{datetime.now().strftime("%Y年%m月%d日")}', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, 20, 0, 1.5, 0)

# ==================== 摘要 ====================
page_break()
h1('摘  要', 1)

body('随着全球化供应链日益复杂和社交媒体信息传播速度加快，企业面临的舆情风险、供应链中断风险和法律合规风险呈上升趋势。传统的"监控—人工判断—人工决策"模式存在响应滞后、信息碎片化和缺乏量化评估三大痛点。本文设计并实现了一个基于编排式Agent架构的企业风险监控智能体系统——"风险哨兵"（Risk Sentinel），以IF椰子水品牌为应用案例，将多源数据聚合、NLP情感分析、定量风险评估（EL/VaR）和多维应对策略生成整合为一个可自动运行的Web系统。系统由21个Python模块（7,089行）和5个前端模板（4,945行）构成，包含16个API端点。核心技术创新包括：三阶段编排流水线（初筛→深度分析→动态预警）、将信用风险EL模型迁移到舆情场景、五维风险应对策略库（供应链/品牌/法律/市场/财务）和规则引擎驱动的优先级排序算法。系统测试表明：新闻管道可用率90%，EL模型能反映财务脆弱性对风险的放大效应（调整前后EL从12.67升至41.04），应对引擎正确识别"三重风险叠加"局面并输出优先级排序的应对方案。')

p = body_ni(''); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.add_run('关键词：').bold = True
p.add_run('Agent编排；风险监控；NLP情感分析；预期损失（EL）；策略库；IF椰子水')

# ==================== 二、引言 ====================
page_break()
h1('一、引  言', 2)

h2('1.1 项目背景')

body('在当今全球化商业环境中，快消品（FMCG）企业面临的风险来源日益多元化——从产品质量质疑、供应链中断、到监管处罚和竞品趁虚而入，风险传导的速度远超传统风险管理体系的响应能力。以IF椰子水为例，该品牌是源自泰国的进口椰子水品牌，在中国市场占据约14%的份额。作为一个高度依赖泰国原材料供应（Nam Hom香水椰100%依赖）、跨境物流、进口食品安全合规和消费者品牌信任的快消品，IF椰子水面临着供应链风险、法律合规风险、品牌声誉风险和市场竞争风险的多重交织。')

body('传统风险管理依赖季度报告和人工审计，存在三大痛点：（1）信息滞后——从风险事件发生到管理层获知可能需要数天甚至数周；（2）数据孤岛——新闻舆情、电商数据、供应链信息、企业工商数据各自分散，无法形成统一的风险视图；（3）缺乏量化——风险评估多依赖定性判断，缺少可量化的风险暴露值计算和财务影响预估。')

h2('1.2 项目目标')

body('本项目的目标是构建一个可自动运行的Web应用系统（Agent），实现对IF椰子水品牌的多维度风险监控。具体目标包括：（1）自动采集和聚合多源风险相关数据，包括新闻舆情、电商数据、企业工商信息、股价财务数据和政策法规；（2）基于NLP和定量模型自动完成风险识别、风险量化和风险等级判定；（3）根据风险信号自动生成优先级排序的应对方案和执行时间线；（4）提供实时预警和可视化仪表盘展示。')

h2('1.3 技术路线')

body('系统采用Python Flask作为Web框架，前端使用Chart.js实现数据可视化。核心技术路线包括：使用策略模式管理多类型新闻源（NewsAPI/GNews/RSS/Demo），通过Baidu NLP + TextBlob + SnowNLP三级回退实现中英文情感分析路由，基于EL（预期损失）= PD × LGD × EAD模型进行定量风险评估，采用playbook模式加载五维应对策略库，通过三阶段编排流水线实现从数据采集到应对方案生成的端到端自动化。')

# ==================== 三、需求分析 ====================
page_break()
h1('二、需求分析', 3)

h2('2.1 功能需求')

add_table(
    ['功能模块', '功能点', '优先级', '描述'],
    [
        ['多源新闻聚合', 'API+RSS多源抓取', 'P0', '从NewsAPI、GNews、7个RSS源获取新闻，300秒超时，智能去重'],
        ['情感分析', '多语言情感分析', 'P0', '中文→Baidu NLP→SnowNLP→TextBlob三级回退'],
        ['风险识别', '风险矩阵评估', 'P0', '计算概率+影响度→高/中/低等级+评分(0-100)'],
        ['定量风险评估', 'EL/VaR/压力测试/情景分析', 'P1', 'PD×LGD×EAD模型+四情景概率加权'],
        ['智能体编排', '三阶段工作流', 'P0', '初筛→条件深度分析→结构化报告'],
        ['风险应对引擎', '五维策略库+优先级排序', 'P1', '供应链/品牌/法律/市场/财务+时间线+成本估算'],
        ['实时预警', '负面情感>0.6+预警关键词', 'P1', '控制台+存储+前端红绿灯仪表盘'],
        ['扩展监控', '电商/企业/财务/政策/卡脖子', 'P2', '5个独立monitor模块'],
        ['报告生成', '自动周报', 'P2', '7模块聚合+红绿灯+环比+TOP3风险'],
    ],
    '系统功能需求矩阵',
    [2.5, 3.5, 1.0, 7.0]
)

h2('2.2 非功能需求')

body('（1）可用性：系统应能在国内网络环境下正常运行，API不可用时自动降级为Demo数据源。（2）性能：新闻抓取300秒超时，缓存命中后<3秒响应。（3）可扩展性：新增监控维度只需注册新数据源模块和策略模板。（4）可维护性：所有配置集中在config.py，模块间通过松耦合的API通信。')

h2('2.3 用户角色与用例')

add_table(
    ['用户角色', '职责', '主要用例', '使用频率'],
    [
        ['风险管理师', '日常监控+风险研判', '查看仪表盘/运行分析/生成周报', '每日'],
        ['品牌公关负责人', '舆情应对', '查看风险热图/获取应对方案/查看预警', '事件驱动'],
        ['供应链经理', '供应链风险监控', '查看卡脖子分析/供应商切换成本评估', '每周'],
        ['合规官', '法规红线监控', '查看政策红线面板/合规缺口追踪', '每周'],
        ['高管/CEO', '全局态势感知', '查看周报摘要/红绿灯仪表盘/TOP3风险', '每周'],
    ],
    '用户角色与用例',
    [2.5, 3.0, 5.5, 2.0]
)

h2('2.4 竞品分析')

add_table(
    ['维度', '传统风控SaaS\n(Meltwater/Cision)', '通用AI Agent\n(ChatGPT+插件)', '本系统\n(Risk Sentinel)'],
    [
        ['数据聚合', '新闻+社交媒体(付费API)', '依赖用户手动输入', '新闻+电商+工商+财务+法规(免费)'],
        ['风险量化', '关键词计数+趋势图', 'LLM推理(不可复现)', 'EL/VaR/压力测试/四情景(可审计)'],
        ['应对方案', '无(仅数据展示)', 'LLM生成(可能幻觉)', '策略库+规则引擎(确定性输出)'],
        ['成本', '数千-数万美元/年', '按token计费(不可控)', '零API调用费(纯本地)'],
        ['可定制性', '低(SaaS标准化)', '中(依赖prompt)', '高(开源+模块化)'],
        ['适用场景', '大型企业', '通用场景', '中小企业垂直风险管理'],
    ],
    '竞品功能对比',
    [2.0, 3.5, 4.0, 4.5]
)

# ==================== 四、系统设计 ====================
page_break()
h1('三、系统设计', 4)

h2('3.1 总体架构')

body('系统采用"数据采集—分析引擎—应对决策—前端展示"四层架构。数据采集层使用策略模式管理多类型新闻源（API/RSS/Demo），BrandFilter和NewsCache为横切关注点。分析引擎层各模块零耦合——risk_analyzer负责NLP+风险矩阵，quantitative_risk负责EL/VaR/情景分析，baidu_nlp负责多语言情感路由。应对决策层通过playbook模式加载五维策略模板，输出优先级排序的行动项。前端展示层通过10个异步面板按依赖关系依次加载渲染。')

h2('3.2 Agent角色与协作流程')

body('系统采用编排式Agent架构，定义了以下Agent角色及其协作关系：')

add_table(
    ['Agent角色', '职责', '输入', '输出', '触发条件'],
    [
        ['数据采集Agent', '多源新闻聚合+去重+缓存', 'RSS源/GNews/Demo', '标准化新闻列表', '定时/用户触发'],
        ['NLP分析Agent', '情感分析+风险关键词匹配', '新闻文本', '情感极性+风险类型', '数据采集完成'],
        ['风险评估Agent', '风险矩阵+EL/VaR+情景分析', 'NLP结果+财务数据', '风险等级+EL值', 'NLP分析完成'],
        ['预警Agent', '实时关键词+情感阈值检测', '逐条新闻分析结果', '预警通知', '每条新闻分析时'],
        ['应对决策Agent', '策略库匹配+优先级排序', '全维度风险信号', '行动项+时间线+成本', '风险评估完成'],
        ['报告生成Agent', '周报聚合+红绿灯+对比', '7模块快照+历史', '结构化周报', '用户触发/每周'],
    ],
    'Agent角色定义与协作关系',
    [2.0, 2.5, 2.5, 3.5, 2.5]
)

h2('3.3 核心模块说明')

body('news_fetcher.py（894行）：多源新闻聚合模块，采用策略模式+抽象基类设计，支持NewsAPI、GNews、RSS和Demo四种源类型。通过BrandFilter四级过滤（品牌关键词→正则精确→品类→行业源）确保相关性，NewsCache 4小时TTL缓存。')

body('risk_analyzer.py（498行）：NLP风险分析模块。SentimentAnalyzer实现中英文智能路由（Baidu NLP→SnowNLP→TextBlob三级回退）。identify_risks()使用正则匹配6大风险类别（劳工/法律/供应链/财务/声誉/环境）的中英双语关键词。calculate_probability()综合来源权威性、关键词强度和时效性三个维度计算风险概率（0-1）。')

body('quantitative_risk.py（550行）：定量风险评估模块。将金融EL模型迁移到舆情场景，新增财务比率调整因子（6维度加权→0.9-1.8×乘数）和四情景概率加权分析（乐观15%/基准50%/悲观25%/极端10%）。')

add_equation('EL = PD × LGD × EAD', ctr.next_e())
add_equation('PD = min(NegativeRatio × AvgIntensity × θ_fin, 0.5)', ctr.next_e())
add_equation('Weighted_EL = Σ(Scenario_EL_i × Probability_i)', ctr.next_e())

body('risk_response.py（792行）：风险应对引擎。5个维度的应对策略各自封装为独立playbook，行动优先级 = 紧急度×影响力×危机加权（RED×1.5）。输出三阶段执行时间线（48h/2周/3月）。')

# ==================== 五、AI技术与Agent实现 ====================
page_break()
h1('四、AI技术与Agent实现', 5)

h2('4.1 模型与工具选型')

add_table(
    ['技术组件', '选型', '选型理由', '备选方案'],
    [
        ['NLP-英文情感', 'TextBlob (NLTK)', '离线可用、无API限制、足够应对新闻文本', 'VADER, BERT fine-tune'],
        ['NLP-中文情感', 'Baidu NLP API (主)\nSnowNLP (备)', '百度5万次/天免费、中文精度高\nSnowNLP离线兜底', 'HanLP, 讯飞API'],
        ['量化计算', 'NumPy', '百分位数/VaR/矩阵运算', 'SciPy, Pandas'],
        ['规则引擎', 'Python dict + 策略库', '轻量、透明、可审计，无需引入规则引擎框架', 'Drools, JSON Rules'],
        ['编排框架', '自研三阶段流水线', '领域专用、零依赖、调试透明', 'LangChain, AutoGen'],
        ['前端', 'Chart.js + CSS Grid', '轻量CDN、无需构建工具、动态异步渲染', 'ECharts, D3.js'],
        ['股价数据', 'yfinance', '免费、支持港股(6603.HK)、自动缓存', 'akshare, 东方财富API'],
    ],
    '模型与工具选型',
    [2.5, 3.0, 5.5, 3.0]
)

h2('4.2 Prompt/规则引擎设计')

body('本系统采用"编排式Agent"架构，核心决策逻辑基于规则引擎+策略库（playbook），而非LLM prompt。这一设计选择基于以下考量：（1）风险类型明确——IF椰子水面临的风险可穷举为供应链/法律/品牌/财务/市场五类；（2）决策规则可穷举——每类风险都有对应的行业最佳实践和标准应对流程；（3）审计要求——金融和合规领域要求决策路径可追溯、可解释。')

body('规则引擎的核心数据结构是五维策略库（playbook），每个策略库包含：触发条件（triggers，如supply_chain_risk、negative_review_surge）、响应模板（response_templates，含具体行动项、成本估算、时间线）和优先级权重（urgency 1-10 × impact 1-10）。当系统检测到触发条件满足时，引擎自动匹配对应的响应模板，计算优先级评分，并按降序输出行动项列表。')

body('以品牌危机应对策略库为例：当微博负面声量>50%且天猫评分<4.5时，触发"红色预警"级别响应，匹配"公开声明+CEO回应"行动模板，优先级评分=10(紧急度)×9(影响力)×1.5(危机加权)=135，排序为最高优先级。')

h2('4.3 LLM集成的设计预留')

body('当前系统不使用LLM进行核心决策，但在以下场景预留了LLM集成接口：（1）声明生成——在品牌危机应对的"声明策略"环节，可将策略类型（如acknowledge/clarify）和上下文输入LLM，自动生成符合品牌调性的公关声明草稿；（2）消费者投诉归纳——将差评关键词输入LLM进行主题聚类和摘要生成；（3）What-If对话——用户输入自然语言问题（如"如果泰国洪水，对IF有何影响"），LLM解析意图后调用情景分析引擎重算EL并返回自然语言分析。这些集成点遵循"规则引擎选策略，LLM写内容"的分工原则——用规则保证下限（可控、可审计），用LLM提升上限（灵活、自然）。')

# ==================== 六、系统实现与演示 ====================
page_break()
h1('五、系统实现与演示', 6)

h2('5.1 核心功能实现')

body('系统实现了7个独立的监控维度，每个维度对应一个前端面板和API端点。核心数据流为：多源新闻→NewsFetcher（去重+缓存+品牌过滤）→RiskAnalyzer（情感+风险矩阵）→QuantitativeRisk（EL/VaR/情景）→RiskResponse（策略库+优先级排序）→JSON→前端异步渲染。')

h2('5.2 以IF椰子水为案例的演示')

body('以下展示系统运行效果。启动python app.py后，浏览器访问http://localhost:5000，系统自动执行三阶段编排分析，页面依次加载10个面板。')

add_figure('report_charts/chart_traffic_light.png', '七模块红绿灯仪表盘：新闻舆情/定量风险/电商/企业/财务/政策/卡脖子')
add_figure('report_charts/chart_news_pipeline.png', '新闻抓取管道：109条原始→去重86→品牌过滤49→+Demo 25→71条入分析')
add_figure('report_charts/chart_scenarios.png', '四情景概率加权EL：乐观(15%)→基准(50%)→悲观(25%)→极端(10%)')

h2('5.3 前端演示截图指引')

body('建议截取以下位置的前端页面：（1）首页全景——所有面板加载完成后的完整页面；（2）风险热图——"🔥 风险热图"面板的3×3矩阵；（3）定量分析——EL/PD/LGD/VaR指标卡；（4）风险应对方案——"🛡️ 风险应对方案"面板的TOP5优先行动排序表；（5）红绿灯仪表盘——"🚦 风险态势红绿灯仪表盘"面板。')

# ==================== 七、评估与讨论 ====================
page_break()
h1('六、评估与讨论', 7)

h2('6.1 系统 vs 传统人工方法')

add_table(
    ['评估维度', '传统人工方法', 'Risk Sentinel系统', '提升幅度'],
    [
        ['数据采集速度', '数小时至数天（人工浏览+整理）', '<5分钟（自动化多源聚合+去重）', '>100倍'],
        ['风险识别覆盖率', '依赖个人经验，可能遗漏', '6风险类别×36关键词×双语', '系统性覆盖'],
        ['风险量化能力', '定性判断为主', 'EL/VaR/压力测试/四情景', '从定性到定量'],
        ['应对方案生成', '人工讨论+手工制定', '策略库自动匹配+优先级排序', '从小时到秒级'],
        ['监测持续性', '工作日8小时', '7×24小时（可扩展为定时轮询）', '3倍时间覆盖'],
        ['成本', '1-2名全职风险管理师', '服务器电费+免费API', '接近零边际成本'],
        ['可审计性', '依赖个人笔记和邮件', '全流程日志+快照+JSON存储', '完整审计追踪'],
    ],
    '系统与人工方法对比',
    [2.0, 3.0, 4.5, 4.5]
)

h2('6.2 局限性分析')

body('（1）数据真实性：电商和社交媒体数据主要基于模拟，真实API的获取受平台反爬限制。在真实企业部署中，需要与平台建立数据合作或使用官方商业数据产品。（2）NLP精度：通用模型未针对椰子水行业术语微调，"sugar added"可能被误判为中性。引入领域微调BERT或LLM few-shot分类可提升。（3）EL模型简化：EAD通过财务调整因子动态计算，但PD上限0.5导致极端情况下风险被低估。更精确的尾部风险建模需要引入极值理论（EVT）。（4）单品牌局限：当前为IF椰子水定制，多品牌SaaS化需要参数化品牌配置和策略库内容。')

h2('6.3 改进空间')

body('（1）实时管道升级——从"请求-响应"模式升级为APScheduler定时轮询+WebSocket推送。（2）供应链数字孪生——接入天气API+海运AIS+椰子期货价格。（3）情绪速率追踪——增加时间窗口计算，捕捉"1小时内负面飙升15%"的早期信号。（4）多品牌SaaS化——将品牌特定配置参数化，提供管理后台。')

# ==================== 八、总结 ====================
h1('七、总  结', 8)

body('本文设计并实现了基于编排式Agent架构的企业风险监控智能体——"风险哨兵"，以IF椰子水品牌为应用案例完成了端到端验证。系统包含21个Python模块（7,089行）和5个前端模板（4,945行），涵盖16个API端点和7个独立监控维度。核心贡献包括：（1）提出了三阶段编排流水线，根据初筛结果动态决定分析深度；（2）将信用风险EL模型成功迁移到舆情场景，并引入财务比率调整因子；（3）构建了五维风险应对策略库，实现从风险信号到可执行建议的自动化闭环；（4）验证了"编排式Agent"在垂直风险领域的工程实用性——比LLM Agent更可控、更便宜、更可审计。')

body('系统的设计哲学可概括为"小而聚焦"——不试图构建一个能回答任何风险问题的通用AI，而是构建一个在IF椰子水的风险域内足够专业、足够快、足够便宜的专用Agent。这种模式特别适合中小企业风险管理场景——数据量有限但风险类型明确，需要的是确定性的决策支持而非创造性的发散推理。')

# ==================== 附录A: AI使用说明 ====================
page_break()
h1('附  录', 9)
h2('附录A  AI使用说明')

body('本系统在以下环节使用了AI/NLP技术：')

add_table(
    ['使用环节', 'AI技术/工具', '使用方式', '说明'],
    [
        ['英文情感分析', 'TextBlob (NLTK)', '离线调用', '对英文新闻文本进行情感极性分析，输出-1~1极性分数'],
        ['中文情感分析', 'Baidu NLP API', '在线API调用', '通过OAuth 2.0认证，调用百度自然语言处理-情感分析接口'],
        ['中文情感(离线备)', 'SnowNLP', '离线调用', '百度API不可用时的中文情感分析降级方案'],
        ['风险关键词匹配', '正则表达式+词典', '规则引擎', '基于预设的6类中英双语风险关键词库进行文本匹配'],
        ['新闻去重', 'difflib.SequenceMatcher', '离线调用', '计算标题文本相似度（阈值80%）进行模糊去重'],
        ['股价数据', 'yfinance (Yahoo Finance)', '在线API调用', '获取IFBH(6603.HK)实时股价和财务指标'],
        ['数值计算', 'NumPy', '离线调用', 'VaR百分位数计算、情景分析矩阵运算'],
    ],
    'AI/NLP技术使用清单',
    [2.0, 3.0, 2.5, 6.5]
)

body('声明：本系统的核心决策逻辑（风险等级判定、应对方案选择、优先级排序）完全基于规则引擎和预设策略库，不依赖LLM（大语言模型）进行自主决策。LLM仅在设计预留接口中存在，当前版本未实际使用。所有分析结果均可追溯到具体的触发信号和决策规则，确保可审计性。')

# ==================== 附录B: 参考资料 ====================
h2('附录B  参考资料与开源项目')

body('本系统在开发过程中参考了以下资料并使用了以下开源技术：')

body_ni('一、学术/行业参考资料：')
body('（1）COSO《企业风险管理——整合框架》（2017版）— 提供了ERM八大要素的理论基础。')
body('（2）ISO 31000:2018《风险管理指南》— 风险识别、分析、评价的标准流程参考。')
body('（3）巴塞尔协议III — 信用风险EL模型（EL=PD×LGD×EAD）的原始框架来源。')
body('（4）Jorion, P. 《Value at Risk: The New Benchmark for Managing Financial Risk》— VaR百分位数方法的理论基础。')
body('（5）港交所《上市规则》附录D1A — IFBH招股说明书中的风险因素披露章节。')
body('（6）GB 7718-2025《预包装食品标签通则》— 进口食品标签合规基准。')
body('（7）GB 28050-2025《预包装食品营养标签通则》— 营养标签1+6新规基准。')
body('（8）海关总署第280号令（2025）— 进口食品境外生产企业注册管理规定。')

body_ni('二、开源项目（均已获得原作者许可或使用MIT/BSD等允许商业使用的许可证）：')
body('（9）Flask — Python Web微框架，BSD协议。')
body('（10）TextBlob — 英文情感分析库，MIT协议。')
body('（11）SnowNLP — 中文情感分析库，MIT协议。')
body('（12）feedparser — RSS/Atom解析库，BSD协议。')
body('（13）Chart.js — 前端图表库，MIT协议。')
body('（14）NumPy — 数值计算库，BSD协议。')
body('（15）yfinance — Yahoo Finance数据接口，Apache 2.0协议。')

# ==================== 附录C: 其他补充 ====================
h2('附录C  其他补充说明')

body('一、API端点完整列表（16个）：')
body('GET /（首页）、/api/analyze（核心分析）、/api/health（健康检查）、/api/history（历史记录）、/api/history/{id}（记录详情）、POST /api/history/save（保存）、/api/alerts（预警列表）、POST /api/alerts/clear（清除预警）、POST /api/cache/clear（清除缓存）、/api/ecommerce/monitor（电商监控）、/api/enterprise/monitor（企业风险）、/api/financial/monitor（财务风险）、/api/policy/monitor（政策红线）、/api/response/plan（应对方案）、/api/report/weekly（自动周报）、/api/delphi-analyze（德尔菲评估）。')

body('二、项目代码仓库：本地Git仓库位于 C:/Users/23994/risk_sentinel/，共3次提交，首次提交2026-04-02。核心Python模块21个共7,089行，HTML模板5个共4,945行，总计12,034行代码。')

body('三、数据真实性声明：本系统的电商监控和社交媒体数据采用模拟数据（明确标注），新闻数据混合了真实RSS源（7个）、GNews API数据和DemoSource模拟数据（25条）。IFBH(6603.HK)股价和财务数据通过yfinance从Yahoo Finance获取（免费、公开），政策法规数据来自国家卫健委、海关总署、市场监管总局的公开文件。')

body('四、系统启动方式：在项目根目录运行 python app.py，访问 http://localhost:5000。依赖安装：pip install -r requirements.txt（额外依赖：yfinance, PyPDF2）。')

# ===== 结尾 =====
body_ni('')
add_p('— 报告完 —', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, 30, 10, 1.5, 0)

# ===== 保存 =====
out = r'C:\Users\23994\risk_sentinel\Risk_Sentinel_Final_Report.docx'
doc.save(out)
print(f'Saved: {out}')
print(f'Size: {os.path.getsize(out)} bytes')
