"""
生成Word报告：风险哨兵智能体 Agent开发实践
格式：宋体五号(10.5pt) / 1.5倍行距 / A4纸 / 题注 / 公式编号
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

doc = Document()
CHART_DIR = 'report_charts'

# ===== 页面设置 =====
s = doc.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)
s.top_margin, s.bottom_margin = Cm(3.7), Cm(3.5)
s.left_margin, s.right_margin = Cm(2.8), Cm(2.6)

# ===== 全局计数器 =====
class Counter:
    def __init__(self): self.ch = 1; self.t = 0; self.f = 0; self.eq = 0
    def chapter(self, n): self.ch = int(str(n)[0]) if isinstance(n, str) else n; self.t = 0; self.f = 0; self.eq = 0
    def next_table(self): self.t += 1; return f"{self.ch}-{self.t}"
    def next_fig(self): self.f += 1; return f"{self.ch}-{self.f}"
    def next_eq(self): self.eq += 1; return f"({self.ch}-{self.eq})"

ctr = Counter()

# ===== 样式工具 =====
def fr(run, cn, en, sz, bold=False, color=None):
    run.font.size = Pt(sz); run.bold = bold; run.font.name = en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    if color: run.font.color.rgb = RGBColor(*color)

def add_p(text, cn='宋体', en='Times New Roman', sz=10.5, bold=False, align=None,
          sb=0, sa=0, ls=1.5, indent=0.74, color=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    pf = p.paragraph_format
    pf.space_before=Pt(sb); pf.space_after=Pt(sa); pf.line_spacing=ls
    if indent: pf.first_line_indent = Cm(indent)
    r = p.add_run(text); fr(r, cn, en, sz, bold, color); return p

def body(text): return add_p(text, '宋体', 'Times New Roman', 10.5, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 1, 1, 1.5, 0.74)
def body_ni(text): return add_p(text, '宋体', 'Times New Roman', 10.5, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 1, 1, 1.5, 0)
def title_p(text): return add_p(text, '黑体', 'Times New Roman', 22, True, WD_ALIGN_PARAGRAPH.CENTER, 30, 16, 1.3, 0)
def subtitle_p(text): return add_p(text, '仿宋', 'Times New Roman', 15, False, WD_ALIGN_PARAGRAPH.CENTER, 6, 6, 1.3, 0)
def h1(text, ch_num=None):
    if ch_num: ctr.chapter(ch_num)
    return add_p(text, '黑体', 'Times New Roman', 16, True, WD_ALIGN_PARAGRAPH.LEFT, 18, 10, 1.3, 0)
def h2(text): return add_p(text, '黑体', 'Times New Roman', 14, True, WD_ALIGN_PARAGRAPH.LEFT, 12, 6, 1.3, 0)
def h3(text): return add_p(text, '宋体', 'Times New Roman', 12, True, WD_ALIGN_PARAGRAPH.LEFT, 8, 4, 1.3, 0)
def cover_line(text, sz=14): return add_p(text, '仿宋', 'Times New Roman', sz, False, WD_ALIGN_PARAGRAPH.CENTER, 3, 3, 2.0, 0)
def page_break(): doc.add_page_break()

# ===== 表格（题注在表格上方）=====
def add_table(headers, rows, caption_text, cw=None):
    num = ctr.next_table()
    # 题注（表格上方）
    add_p(f'表{num}  {caption_text}', '宋体', 'Times New Roman', 9, True, WD_ALIGN_PARAGRAPH.CENTER, 8, 2, 1.1, 0)
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(h)
        fr(r,'黑体','Times New Roman',9,True); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>'))
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; r=c.paragraphs[0].add_run(str(val))
            fr(r,'宋体','Times New Roman',9,False)
            c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if ci>0 else WD_ALIGN_PARAGRAPH.LEFT
    if cw:
        for i,w in enumerate(cw):
            for row in t.rows: row.cells[i].width=Cm(w)
    return t

# ===== 图片（题注在图片下方）=====
def add_figure(img_path, caption_text, width_inches=5.5):
    if os.path.exists(img_path):
        num = ctr.next_fig()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format; pf.space_before=Pt(8); pf.space_after=Pt(2)
        r = p.add_run(); r.add_picture(img_path, width=Inches(width_inches))
        add_p(f'图{num}  {caption_text}', '宋体', 'Times New Roman', 9, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 10, 1.1, 0)
    else:
        body_ni(f'【图片占位：{img_path} —— {caption_text}】')

# ===== 公式 =====
def add_equation(equation_text, eq_num=None):
    """公式独占一行居中，右侧编号"""
    if eq_num is None: eq_num = ctr.next_eq()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before=Pt(8); pf.space_after=Pt(8); pf.line_spacing=1.3
    r = p.add_run(equation_text); fr(r, 'Times New Roman', 'Times New Roman', 11, False)
    # 编号
    r2 = p.add_run(f'    {eq_num}'); fr(r2, 'Times New Roman', 'Times New Roman', 10, False)
    return p

# ===== 页眉页脚 =====
for sec in doc.sections:
    hdr=sec.header; hdr.is_linked_to_previous=False
    hp=hdr.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    hr=hp.add_run('风险哨兵智能体——从舆情监控到智能风险决策的Agent开发实践')
    fr(hr,'宋体','Times New Roman',9,False)
    ftr=sec.footer; ftr.is_linked_to_previous=False
    fp=ftr.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fr1=fp.add_run('— '); fr(fr1,'宋体','Times New Roman',9,False)
    fp._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    fp._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    fp._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    fr2=fp.add_run(' —'); fr(fr2,'宋体','Times New Roman',9,False)

# ==================== 封面 ====================
for _ in range(8): cover_line('')
title_p('风险哨兵智能体')
subtitle_p('——从舆情监控到智能风险决策的Agent开发实践')
cover_line('')
subtitle_p('基于编排式Agent架构的企业风险管理平台')
subtitle_p('以IF椰子水品牌为应用案例')
cover_line('')
cover_line('技术栈：Python / Flask / NLP / Chart.js / yfinance', 12)
cover_line('代码规模：21个模块 · 12,000+行代码 · 16个API端点', 12)
cover_line('开发周期：2026年4月至6月', 12)
cover_line('')
cover_line(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}', 12)

# ==================== 一、摘要 ====================
page_break()
h1('一、摘  要', 1)

body('在企业风险管理领域，传统的"监控—人工判断—人工决策"模式面临着响应滞后、信息碎片化和缺乏量化评估三大痛点。本文介绍了一个基于编排式Agent架构的企业风险监控智能体——"风险哨兵"（Risk Sentinel）的设计与实现过程。该智能体以IF椰子水品牌为应用案例，通过"感知—推理—行动"的Agent闭环，将多源数据聚合、NLP情感分析、定量风险评估（EL/VaR）和多维应对策略生成整合为一个可自动运行的Web系统。')

body('系统的核心创新在于采用了"编排式Agent"而非"自主式LLM Agent"的架构选择——通过三阶段编排流水线（初筛→深度分析→动态预警）、五维策略库（供应链/品牌/法律/市场/财务）和规则引擎驱动的优先级排序算法，在可控性、可解释性和成本可控的前提下，实现了从风险信号检测到结构化应对方案的端到端自动化。')

body('系统由21个Python模块（7,089行）和5个前端模板（4,945行）构成，总计约12,000行代码，经过四个阶段的迭代开发，已从基础舆情检测演进为具备电商监控、企业风险、实时股价、政策红线、卡脖子分析、多情景模拟和自动周报生成的综合商业情报与风险管理平台。')

p = body_ni('')
r1 = p.add_run('关键词：')
fr(r1, '黑体', 'Times New Roman', 10.5, True)
r2 = p.add_run('Agent编排；风险监控；NLP情感分析；预期损失（EL）；策略库；智能体；IF椰子水')
fr(r2, '宋体', 'Times New Roman', 10.5, False)

# ==================== 二、智能体设计理念 ====================
page_break()
h1('二、智能体设计理念', 2)

h2('2.1 为什么是Agent而不是传统SaaS')

body('传统的企业风险监控SaaS工具通常遵循"仪表盘→人工判断→人工决策"的单向信息流：系统负责数据采集和可视化展示，风险管理专家查看仪表盘后，依靠个人经验做出判断并手动执行响应措施。这种模式的局限性在IF椰子水案例中尤为突出——当品牌同时面临舆情危机（微博负面55%）、财务恶化（增收不增利）、合规漏洞（标签违规）和供应链瓶颈（100%依赖单一采购商）的多维风险叠加时，传统仪表盘只能呈现各维度的独立指标，无法形成统一的"风险态势认知"，更无法自动生成跨维度的协同应对方案。')

body('Agent模式的核心区别在于引入了一个"推理与行动"闭环。在风险哨兵智能体中，这个闭环通过三阶段编排流水线实现——感知层聚合7个数据源，推理层使用风险矩阵+EL/VaR+情景分析进行多维度量化，行动层通过五维策略库生成优先级排序的应对方案。')

h2('2.2 编排式Agent vs 自主式Agent')

body('在架构选择上，我们面临一个关键决策：使用基于大语言模型（LLM）的自主式Agent，还是基于规则引擎+策略库的编排式Agent？')

add_table(
    ['维度', '编排式Agent（本项目选择）', '自主式LLM Agent'],
    [
        ['可控性', '高——所有决策路径可追溯、可审计', '低——LLM的黑箱推理不可预测'],
        ['成本', '极低——纯CPU计算，无API调用费', '高——每次推理需调用LLM API'],
        ['可解释性', '强——每个建议可追溯到触发信号和策略模板', '弱——LLM的输出难以逐条解释'],
        ['领域精度', '高——策略库由领域专家预设，确定性输出', '中——依赖prompt质量，可能产生幻觉'],
        ['灵活性', '中——新增领域需扩展策略库', '高——可通过prompt快速适配新场景'],
        ['适用场景', '风险类型明确、规则可穷举的中小企业', '高度不确定、需创造性推理的复杂场景'],
    ],
    '编排式Agent与自主式LLM Agent的比较',
    [2.0, 6.0, 6.0]
)

body('对于IF椰子水这样的中小企业风险管理场景——风险类型明确（供应链、法律、品牌、财务、市场五类）、决策规则可穷举（每种风险都有对应的行业最佳实践）——编排式Agent是更务实的选择。')

h2('2.3 三阶段编排流水线')

body('智能体编排流水线是本系统的核心设计。与传统的一次性全量分析不同，编排工作流根据初步分析结果动态决定后续分析深度。')

body('第一步（快速初筛）：对71篇文章逐条进行情感分析+风险矩阵评估+实时预警检查，输出高/中/低风险分类。第二步（深度分析——条件触发）：仅当中高风险文章>0时才执行EL/VaR/压力测试/四情景分析，低风险时跳过以节省约40%的分析时间。第三步（报告生成）：聚合所有维度的结果，输出JSON给前端10个面板异步渲染。')

add_figure('report_charts/chart_code_growth.png', '智能体编排三阶段流水线（前端截图建议：点击页面"开始分析"后，观察控制台输出的三阶段日志）', 5.0)

add_figure('report_charts/chart_news_pipeline.png', '新闻抓取管道数据流——从10个原始数据源到71条分析输入', 5.0)

# ==================== 三、技术架构 ====================
page_break()
h1('三、技术架构', 3)

h2('3.1 技术栈总览')

add_table(
    ['层级', '技术选型', '用途说明'],
    [
        ['Web框架', 'Flask 3.0', 'REST API服务，16个端点，支持本地+云端部署'],
        ['前端渲染', 'Chart.js + CSS Grid', '赛博朋克风格SPA，10个面板异步加载'],
        ['英文NLP', 'TextBlob (NLTK)', '英文情感分析（离线，无API限制）'],
        ['中文NLP', 'Baidu NLP + SnowNLP', '中文情感分析，百度API(5万次/天)+离线兜底'],
        ['量化计算', 'NumPy', 'VaR百分位数、情景分析矩阵'],
        ['股价数据', 'yfinance', 'IFBH(6603.HK)实时股价，回退港交所缓存'],
        ['RSS解析', 'feedparser', '7个RSS源的XML内容解析'],
        ['去重引擎', 'difflib.SequenceMatcher', '标题相似度模糊去重（阈值80%）'],
        ['持久化', 'JSON + SQLite', '历史记录（JSON）+ MCP数据仓库（SQLite）'],
    ],
    '系统技术栈',
    [2.5, 4.5, 7.0]
)

h2('3.2 四层架构设计')

body('系统采用"数据采集—分析引擎—应对决策—前端展示"四层架构。数据采集层使用策略模式管理多类型新闻源，分析引擎层各模块零耦合，应对决策层通过playbook模式加载策略模板，展示层通过10个异步面板按依赖关系依次渲染。')

add_figure('report_charts/chart_modules.png', '核心模块代码量分布——Top 15模块覆盖系统87%的代码量（前端截图建议：页面首页整体截图，展示四层架构在UI上的体现）', 5.0)

h2('3.3 核心公式——预期损失（EL）模型')

body('系统将金融信用风险领域的EL模型迁移到舆情风险场景。以下是核心公式：')

add_equation('EL = PD × LGD × EAD', '(3-1)')
body('其中，PD（违约概率的舆情适配）= 负面新闻占比 × 平均负面情感强度，上限为0.5。')

add_equation('PD = min(NegativeRatio × AvgNegativeIntensity × θ_fin, 0.5)', '(3-2)')
body('其中，θ_fin为财务比率调整因子，取值范围0.9-1.8，由6个财务维度的加权评分映射得到。')

add_equation('LGD = Σ(RiskTypeWeight_i) / RiskCount', '(3-3)')
body('风险类型权重预设为：法律风险/供应链风险=0.8，财务风险=0.5，声誉/劳工/环境风险=0.3。')

add_equation('EAD = 100 × θ_fin,  (EAD ≥ 60)', '(3-4)')
body('EAD（风险暴露）不再固定为100，而是随财务调整因子动态变化。财务脆弱的企业在相同舆情冲击下，品牌价值损失更大。')

# ==================== 四、核心功能实现 ====================
page_break()
h1('四、核心功能实现', 4)

h2('4.1 多源数据聚合')

body('新闻获取是整个系统的数据入口，也是最脆弱的环节。初始配置的32个RSS源中仅7个（22%）可在中国大陆网络环境下正常访问。为应对这一挑战，设计了三层冗余架构：第一层——7个经验证RSS源加GNews API提供真实行业新闻；第二层——DemoSource始终运行，提供25条IF椰子水专项模拟数据保底；第三层——BrandFilter从严格过滤调整为分层过滤，使可用新闻量提升3倍。')

body('缓存层（NewsCache）使用JSON文件以源名称为键，4小时TTL。第二次访问直接读取缓存，将响应时间从分钟级降至秒级。去重采用两阶段策略：URL精确匹配→标题相似度模糊去重（difflib.SequenceMatcher，阈值80%）。')

h2('4.2 情感分析与风险矩阵')

body('情感分析模块实现了多语言智能路由：检测到中文字符→Baidu NLP API→失败回退SnowNLP→最后回退TextBlob；纯英文→直接TextBlob。这种三级回退机制确保在任何API可用性条件下都有可用的分析引擎。')

body('风险矩阵由两个函数构成。calculate_probability()对每条新闻计算风险发生概率（0-1），综合三个维度：来源权威性（Reuters/Bloomberg +0.20）、关键词强度（召回/诉讼类 +0.30）、时效性（3天内 +0.10），基础概率0.30，上限0.95。risk_matrix(impact, probability)将影响度（risk_score/100归一化）和概率映射为三级风险等级：高风险（impact>0.5且prob>0.6）、中风险（impact>0.3且prob>0.4或乘积>0.2）、低风险（其余）。')

add_figure('report_charts/chart_traffic_light.png', '七模块风险红绿灯仪表盘——聚合所有监控维度的当前风险状态（前端截图建议：页面加载完成后，滚动到"红绿灯仪表盘"区域截图）', 5.0)

h2('4.3 定量风险评估')

body('EL模型是本系统最具技术特色的创新。下表展示EL三因子的舆情场景适配结果：')

add_table(
    ['因子', '信用风险原定义', '舆情风险适配', 'IFBH当前值', '计算方式'],
    [
        ['PD', '借款人违约概率', '风险事件发生概率', '0.20(raw)→0.36(adj)', '负面占比×平均强度×θ_fin, 上限0.5'],
        ['LGD', '违约损失率', '风险类型的潜在损失程度', '0.63', '法律/供应链=0.8, 财务=0.5, 声誉=0.3'],
        ['EAD', '违约风险暴露', '企业声誉价值基准分', '100(raw)→180(adj)', '100 × θ_fin, 下限60'],
    ],
    'EL模型三因子的舆情风险场景适配',
    [1.2, 3.0, 3.5, 3.0, 3.3]
)

add_figure('report_charts/chart_el_comparison.png', 'EL模型：财务调整因子对EL/PD/EAD的影响对比（无调整基准 vs 1.8x乘数 vs 0.9x折扣）', 5.0)

h2('4.4 多情景分析引擎')

body('将单一EL扩展为四情景概率加权模型，反映不确定性下的期望损失范围。')

add_figure('report_charts/chart_scenarios.png', '四情景概率加权EL——乐观(15%)→基准(50%)→悲观(25%)→极端(10%)', 5.0)

body('加权EL公式：')

add_equation('Weighted_EL = Σ(Scenario_EL_i × Probability_i)', '(4-1)')

body('IFBH当前加权EL为45.2（HIGH），乐观-悲观情景EL差距为35.0（基准的79.6%），反映较高的不确定性。')

h2('4.5 五维风险应对引擎')

body('risk_response.py（792行）是系统第二大模块。核心设计是策略库（Playbook）模式——每个维度的应对策略各自封装为独立的策略模板，包含触发条件、应对行动、成本估算和执行时间线。')

add_figure('report_charts/chart_strategy.png', '五维风险应对策略库——行动项数量与预估成本概览（前端截图建议：滚动至页面底部"风险应对方案"面板，展示TOP5优先行动排序表）', 5.0)

body('行动优先级排序算法：每个行动项的紧急度（1-10）×影响力（1-10）得到基础优先级分，再按危机等级加权（RED×1.5, ORANGE×1.2）。最终按优先级降序排列，分配到三阶段执行时间线（紧急响应48h/短期整改2周/中期巩固3月），每个行动项标注责任部门和成本。')

h2('4.6 扩展监控维度')

body('系统在第四次迭代中，逐步补全了5个独立的监控维度。')

add_table(
    ['模块', '框架对齐', '数据源', 'IFBH当前关键指标'],
    [
        ['电商监控', 'PDF M5:行业地位', '天猫/京东/小红书/抖音(模拟)', '月销45K箱↓15.8% / 评分4.5 / 差评率11%'],
        ['企业风险', 'PDF M8:ESG合规', '天眼查/裁判文书/海关(模拟)', '2次行政处罚 / 3起诉讼 / 通关+0.8天'],
        ['财务监控', 'PDF M7:反常信号', 'yfinance + 港交所财报', '股价HK$13.3(-52% vs IPO) / 6个反常信号'],
        ['政策红线', 'PDF M3:政策红绿灯', '卫健委/海关/市监(真实法规)', '8项法规 / 5个合规缺口 / 1项已逾期'],
        ['卡脖子分析', 'PDF M6:技术壁垒', 'IFBH招股书+供应链实调', '100%单品种 / 100%单采购商 / 100%代工'],
    ],
    '扩展监控维度一览',
    [2.0, 2.5, 4.0, 5.5]
)

# ==================== 五、开发过程演进 ====================
page_break()
h1('五、开发过程演进', 5)

h2('5.1 四个开发阶段')

add_table(
    ['阶段', '时间', '代表模块', '代码量', '核心能力跃迁'],
    [
        ['Phase 1: 基础舆情', '4月初', 'app+risk_analyzer+news_fetcher', '~450行', '"能跑通"：新闻→情感→评分→展示'],
        ['Phase 2: 量化风险', '4月中', '+quantitative_risk+baidu_nlp', '~1,050行', '"可量化"：EL/VaR/压力测试'],
        ['Phase 3: Agent编排', '4月下', '+risk_response+workflow', '~2,150行', '"能应对"：编排+策略库+实时预警'],
        ['Phase 4: 多维扩展', '5-6月', '+5个monitor+情景+周报', '~7,050行', '"面面俱到"：电商/财务/政策/卡脖子'],
    ],
    '系统开发四阶段演进',
    [1.5, 1.3, 4.5, 1.8, 4.9]
)

h2('5.2 关键设计决策')

body('决策一——DemoSource从fallback改为始终运行：初始仅在API全失败时启用，但RSS源对IF品牌的直接覆盖率极低（<5%）。改为始终运行后，可用新闻从不足20条提升至71条。')

body('决策二——BrandFilter从严格改为分层：初始仅保留含"IF"关键词的新闻，损失大量行业背景。调整为分层过滤（饮料行业源保留所有椰子水新闻），新闻量增长3倍。')

body('决策三——EAD从固定值改为动态调整：最初EAD=100，无法区分企业的差异脆弱性。引入6维度财务比率调整因子（0.9-1.8×），EL能反映"财务脆弱企业受同样舆情冲击更大"的现实逻辑。')

h2('5.3 技术挑战与解决方案')

body('挑战一：初始32个RSS源仅7个可用（22%）。解决方案——"三层冗余"（真实API+Demo保底+分层过滤）。')

body('挑战二：Bing/Baidu新闻搜索不返回标准RSS。解决方案——放弃直接RSS解析，通过BrandFilter+行业源间接获取。')

body('挑战三：yfinance在国内网络限流。解决方案——双模式：yfinance优先（实时），失败回退港交所缓存数据。')

body('挑战四：EL模型在极端情况下的PD饱和（上限0.5）。解决方案——通过EAD动态调整（100→180）部分弥补，但更精确的尾部风险建模仍有改进空间。')

# ==================== 六、前端展示与用户体验 ====================
page_break()
h1('六、前端展示与用户体验', 6)

body('前端采用单页面应用（SPA）架构，视觉设计采用赛博朋克风格——深色背景（#0a0a0f）配合青蓝色（#00f5d4）网格线和粉色（#ff006e）点缀色。页面加载后通过10个异步fetch请求依次加载面板：核心分析(1.0s)→历史记录(1.0s)→电商(2.0s)→企业(2.5s)→财务(2.8s)→政策(3.0s)→卡脖子(3.2s)→情景分析(3.3s)→应对方案(3.5s)→周报(3.6s)。')

body('风险热图是系统的核心可视化——CSS Grid实现3×3矩阵，X轴为概率（低/中/高），Y轴为影响度（低/中/高），每个格子动态显示文章计数和颜色渐变（绿→黄→红→粉），本质上是risk_matrix()函数的可视化投影。')

body_ni('【📸 前端截图建议】')
body_ni('')
body_ni('建议在以下6个位置截图插入本报告：')
body_ni('① 首页全景 —— 启动 python app.py 后，等待所有面板加载完成（约4秒），截图整个页面展示10个面板的全貌')
body_ni('② 风险热图区域 —— 滚动到"🔥 风险热图"面板，截图3×3矩阵的细节')
body_ni('③ 定量分析面板 —— 截图EL/PD/LGD/VaR指标卡和压力测试进度条')
body_ni('④ 风险应对方案面板 —— 滚动到底部"🛡️ 风险应对方案"，截图危机等级+优先行动排序表+时间线')
body_ni('⑤ 电商+企业风险面板 —— 截图"🛒 电商监控"和"🏛️ 企业风险监控"两个相邻面板')
body_ni('⑥ 周报面板 —— 截图"📋 本周风险态势报告"，展示红绿灯仪表盘+TOP3风险')
body_ni('')
body_ni('以上6张截图插入位置建议：①插入第四章开头（4.1之前），②插入4.2.4节，③插入4.3节，④插入4.5节，⑤⑥插入4.6节。')

# ==================== 七、测试与验证 ====================
page_break()
h1('七、测试与验证', 7)

h2('7.1 新闻抓取管道测试')

body('测试环境：Windows 11，Python 3.12，国内网络。10个配置源中9个可用（90%），300秒超时内获109条原始→去重86条→品牌过滤49条+Demo 25条=71条入分析管道。缓存命中后响应时间<3秒。')

h2('7.2 定量模型验证')

add_table(
    ['测试场景', 'PD', 'LGD', 'EAD', 'EL', '说明'],
    [
        ['无财务调整（基准EL）', '0.20', '0.63', '100', '12.67', '传统舆情EL'],
        ['融入IFBH财务（θ=1.8×）', '0.36', '0.63', '180', '41.04', '财务脆弱性放大224%'],
        ['财务健康假设（θ=0.9×）', '0.18', '0.63', '90', '10.21', '稳健企业的EL折扣'],
        ['压力测试（负面×1.5）', '—', '—', '—', '+50%', '极端负面爆发冲击'],
        ['加权EL（四情景）', '—', '—', '—', '45.2', '概率加权期望损失'],
    ],
    'EL模型验证结果',
    [3.0, 1.5, 1.5, 2.0, 2.0, 4.0]
)

h2('7.3 应对引擎端到端测试')

body('聚合全维度信号输入RiskResponseEngine，输出：危机等级RED（评分70/100）、10个优先行动项（覆盖5个维度）、TOP1=发布官方声明+CEO回应（优先级135.0）、TOP2=采购去依赖化（优先级135.0）、预估总成本275-315万元、三阶段执行时间线（48h/2周/3月）。引擎正确识别了舆情-供应链-合规"三重风险叠加"的局面。')

# ==================== 八、总结与展望 ====================
page_break()
h1('八、总结与展望', 8)

h2('8.1 项目成果')

body('风险哨兵智能体以12,000行代码的规模，实现了从一个基础舆情检测脚本到一个具备全景商业情报监控与智能风险应对能力的Agent平台的完整演进。系统验证了一种可行的Agent开发范式——在垂直领域内，通过编排式架构+策略库+规则引擎的组合，可以在不依赖LLM的情况下构建出足够"智能"的决策支持系统。')

h2('8.2 "小而聚焦"模式的反思')

body('在风险类型明确、决策规则可穷举的垂直场景中，"编排式Agent"比"自主式LLM Agent"更具工程实用性。这并非对LLM的否定，而是对"何时用规则、何时用模型"的一种务实判断。最自然的LLM集成场景是在风险应对引擎的声明生成环节——"规则引擎选策略，LLM写内容"的分工模式，可能是中小场景Agent应用的最优解。')

h2('8.3 未来方向')

body('（1）LLM增强内容生成——在声明撰写场景集成Claude API，自动生成符合品牌调性的公关声明。（2）实时管道升级——从"请求-响应"升级为APScheduler定时轮询+WebSocket推送的实时架构。（3）多品牌SaaS化——将品牌特定配置参数化，提供管理后台。（4）供应链数字孪生——接入天气API+海运AIS+椰子期货价格。（5）情绪速率追踪——增加时间窗口计算，捕捉"1小时内负面飙升15%"的早期预警信号。')

# ==================== 附录 ====================
page_break()
h1('附  录', 9)

h2('附录A  API端点完整列表')

add_table(
    ['方法', '端点', '功能说明', '优先级'],
    [
        ['GET', '/', '首页渲染（赛博朋克仪表盘）', 'P0'],
        ['GET', '/api/analyze', '核心分析（三阶段编排工作流）', 'P0'],
        ['GET', '/api/health', '健康检查', 'P0'],
        ['GET', '/api/history', '获取最近N条历史记录', 'P1'],
        ['GET', '/api/history/{id}', '单条记录详情', 'P1'],
        ['POST', '/api/history/save', '保存分析结果', 'P1'],
        ['GET', '/api/alerts', '获取活跃实时预警列表', 'P1'],
        ['POST', '/api/alerts/clear', '清除所有预警', 'P1'],
        ['POST', '/api/cache/clear', '强制清除新闻缓存', 'P1'],
        ['GET', '/api/ecommerce/monitor', '电商全平台数据（天猫/京东/小红书等）', 'P2'],
        ['GET', '/api/enterprise/monitor', '企业风险数据（天眼查/裁判文书/海关）', 'P2'],
        ['GET', '/api/financial/monitor', 'IFBH(6603.HK)财务风险报告', 'P2'],
        ['GET', '/api/policy/monitor', '政策红线预警（8项法规+5个缺口）', 'P2'],
        ['GET', '/api/response/plan', '综合风险应对方案生成', 'P1'],
        ['GET', '/api/report/weekly', '自动周报（7模块聚合+环比）', 'P2'],
        ['GET', '/api/delphi-analyze', '德尔菲3轮专家共识评估', 'P2'],
    ],
    'API端点一览（16个）',
    [1.5, 5.0, 5.5, 2.0]
)

h2('附录B  项目文件统计')

add_table(
    ['模块', '行数', '功能定位'],
    [
        ['news_fetcher.py', '894', '数据采集层·多源新闻聚合入口'],
        ['risk_response.py', '791', '应对决策层·五维策略库+优先级排序'],
        ['app.py', '733', '调度层·Flask路由+三阶段编排+预警'],
        ['quantitative_risk.py', '550', '分析引擎层·EL/VaR/压力/情景/财务调整'],
        ['risk_analyzer.py', '498', '分析引擎层·NLP情感+风险矩阵+评分'],
        ['data_sources/ (5个monitor)', '1,665', '扩展层·电商/企业/财务/政策/周报'],
    ],
    '核心Python模块代码统计（完整列表见项目README.md）',
    [5.0, 2.0, 5.0]
)

h2('附录C  外部数据源')

add_table(
    ['数据源', '获取方式', '用途', '限制'],
    [
        ['GNews API', 'HTTP API (api_key)', '英文新闻搜索', '100次/天免费'],
        ['Baidu NLP', 'OAuth 2.0 API', '中文情感分析', '5万次/天免费'],
        ['Yahoo Finance', 'yfinance库', 'IFBH(6603.HK)实时股价', '2000次/小时'],
        ['港交所披露易', '财报PDF', '季度/年度财务数据', '免费公开'],
        ['国家卫健委/海关/市监', '公开法规文件', 'GB标准+进口注册+标签规范', '免费公开'],
        ['RSS Feeds (7个)', 'HTTP+feedparser', '饮料/食品行业新闻背景', '免费'],
        ['TextBlob/SnowNLP', '离线库', '中英文情感分析兜底', '无限制'],
    ],
    '外部API与数据源',
    [2.5, 3.0, 5.0, 3.5]
)

h2('附录D  致谢')

body('本项目使用了Flask（BSD）、TextBlob（MIT）、SnowNLP（MIT）、Chart.js、NumPy、feedparser等开源项目，以及GNews API、Baidu NLP API等免费API服务，在此表示感谢。')

# ===== 结尾 =====
body_ni('')
add_p('— 报告完 —', '宋体', 'Times New Roman', 10.5, False, WD_ALIGN_PARAGRAPH.CENTER, 30, 10, 1.5, 0)

# ===== 保存 =====
out = r'C:\Users\23994\risk_sentinel\Risk_Sentinel_Agent_Report.docx'
doc.save(out)
print(f'Report saved: {out}')
print('File size:', os.path.getsize(out), 'bytes')
