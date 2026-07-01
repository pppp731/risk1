"""
主报告：IFBH Limited (6603.HK) 企业全面风险管理报告
基于公开信息/COSO ERM框架/ISO 31000
格式：五号宋体(10.5pt) / Heading1(黑体16pt) / Heading2(黑体14pt) / 1.5倍行距
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

doc = Document()

# ===== 定义样式 =====
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

h1s = doc.styles['Heading 1']
h1s.font.name = 'Times New Roman'
h1s.font.size = Pt(16); h1s.font.bold = True
h1s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h1s.paragraph_format.space_before = Pt(18)
h1s.paragraph_format.space_after = Pt(10)
h1s.paragraph_format.line_spacing = 1.3

h2s = doc.styles['Heading 2']
h2s.font.name = 'Times New Roman'
h2s.font.size = Pt(14); h2s.font.bold = True
h2s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h2s.paragraph_format.space_before = Pt(12)
h2s.paragraph_format.space_after = Pt(6)
h2s.paragraph_format.line_spacing = 1.3

# ===== 页面设置 =====
s = doc.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)
s.top_margin, s.bottom_margin = Cm(3.7), Cm(3.5)
s.left_margin, s.right_margin = Cm(2.8), Cm(2.6)

# ===== 计数器 =====
class Ctr:
    def __init__(self): self.ch=1; self.t=0; self.f=0
    def c(self,n): self.ch=n; self.t=0; self.f=0
    def nt(self): self.t+=1; return f"{self.ch}-{self.t}"
    def nf(self): self.f+=1; return f"{self.ch}-{self.f}"
ct = Ctr()

# ===== 辅助 =====
def body(text):
    p=doc.add_paragraph(text,style='Normal')
    pf=p.paragraph_format; pf.line_spacing=1.5
    pf.first_line_indent=Cm(0.74); pf.space_before=Pt(1); pf.space_after=Pt(1)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    return p
def body_ni(text):
    p=doc.add_paragraph(text,style='Normal')
    pf=p.paragraph_format; pf.line_spacing=1.5
    pf.space_before=Pt(1); pf.space_after=Pt(1)
    return p
def pb(): doc.add_page_break()

CHART_DIR = 'report_charts'
def add_fig(filename, caption, width=5.0):
    path = os.path.join(CHART_DIR, filename)
    if os.path.exists(path):
        num = ct.nf()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format; pf.space_before = Pt(6); pf.space_after = Pt(2)
        p.add_run().add_picture(path, width=Inches(width))
        cp = doc.add_paragraph(f'图{num}  {caption}')
        cp.style = doc.styles['Normal']; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8); cp.paragraph_format.line_spacing = 1.1
        for r in cp.runs: r.font.size = Pt(9)

def add_table(headers, rows, caption_text, cw=None):
    num=ct.nt()
    p=doc.add_paragraph(f'表{num}  {caption_text}')
    p.style=doc.styles['Normal']; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pf=p.paragraph_format; pf.space_before=Pt(8); pf.space_after=Pt(2); pf.line_spacing=1.1
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in c.paragraphs[0].runs: r.font.size=Pt(9); r.font.bold=True
        c._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>'))
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(val)
            for r in c.paragraphs[0].runs: r.font.size=Pt(9)
    if cw:
        for i,w in enumerate(cw):
            for row in t.rows: row.cells[i].width=Cm(w)
    return t

# ===== 页眉页脚 =====
for sec in doc.sections:
    hdr=sec.header; hdr.is_linked_to_previous=False
    hp=hdr.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    hp.add_run('IFBH Limited (6603.HK)  企业全面风险管理报告')
    hp.runs[0].font.size=Pt(9)
    ftr=sec.footer; ftr.is_linked_to_previous=False
    fp=ftr.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run('— ')
    fp._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    fp._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    fp._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    fp.add_run(' —')

# ==================== 封面 ====================
for _ in range(7): body_ni('')
p=doc.add_paragraph('IFBH Limited (6603.HK)', style='Normal'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph('企业全面风险管理报告', style='Heading 1'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph('—— 基于COSO ERM框架的公开信息分析', style='Normal'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for _ in range(3): body_ni('')
body_ni('分析对象：IFBH Limited（泰国IF椰子水母公司，港交所主板上市）')
body_ni('分析框架：COSO《企业风险管理——整合框架》（2017版）+ ISO 31000:2018')
body_ni('信息来源：港交所披露易（HKEXnews）· IPO招股书 · 2025年年报 · 券商研报 · 公开法规')
body_ni('报告日期：' + datetime.now().strftime('%Y年%m月%d日'))

# ==================== 摘要 ====================
pb()
doc.add_paragraph('摘  要', style='Heading 1')
ct.c(1)

body('IFBH Limited（股票代码：6603.HK）是一家注册于新加坡、运营主体在泰国的饮料企业，旗下拥有"if"和"Innococo"两个品牌，2025年6月30日在香港联交所主板上市。公司以极致的轻资产模式运营——无自有工厂、无仓储设施、仅46名全职员工——通过完全外包的采购、生产和分销体系，构建了年营收1.76亿美元（2025财年）的椰子水业务。公司在中国内地椰子水市场占有率约34%（2024年），连续五年位居第一。')

body('通过系统性的公开信息分析（含IPO招股书风险因素章节、2025年年报、券商研报及公开法规），本报告识别出IFBH面临的五大重大风险：（1）供应链高度集中风险——100%椰子水原料依赖泰国Nam Hom香水椰单一品种，且采购严重依赖控股股东General Beverage；（2）品牌声誉与食品安全风险——2025年多家媒体检测质疑产品质量，社交媒体负面声量达55%；（3）单一市场依赖风险——中国内地市场贡献超过90%的营业收入；（4）财务恶化风险——2025财年增收不增利（营收+11.9%，净利润-31.7%），营销费用飙升76%；（5）法律合规风险——进口食品标签违反GB 28050-2025等新规，已受2次行政处罚。')

body('针对上述风险，本报告提出了包括采购渠道多元化、品牌信任重建、标签合规整改、建立风险准备金等在内的多项应对建议。同时，对公司的治理结构、内控体系和信息披露机制进行了评估，指出关联交易集中、董事会独立性不足、内控资源与业务规模不匹配等改进空间。')

p=body_ni('')
p.add_run('关键词：').bold = True
p.add_run('IFBH；椰子水；企业风险管理；COSO ERM；供应链风险；食品安全；单一市场依赖')

# ==================== 一、公司概况 ====================
pb()
doc.add_paragraph('一、公司概况', style='Heading 1')
ct.c(2)

doc.add_paragraph('1.1 基本信息与股权结构', style='Heading 2')
body('IFBH Limited是一家注册于新加坡共和国的投资控股公司，通过其在泰国的运营实体开展业务。公司于2025年6月30日在香港联交所主板上市（股份代号：06603），IPO发行价为每股27.80港元，募集资金主要用于品牌推广和海外市场拓展。截至2025年12月31日，公司市值为约40亿港元。')

add_table(
    ['项目', '内容', '信息来源'],
    [
        ['公司全称', 'IFBH Limited', '港交所披露易'],
        ['注册地', '新加坡共和国', '2025年年报'],
        ['上市地/日期', '香港联交所主板 / 2025-06-30', '港交所披露易'],
        ['股票代码', '06603.HK', '港交所'],
        ['创始人/CEO', 'Pongsakorn Pongsak（彭萨克）', 'IPO招股书'],
        ['控股股东', 'General Beverage Co., Ltd.（持股约65.51%）', '2025年年报'],
        ['员工人数', '46名全职员工', 'IPO招股书'],
        ['审计师', 'Ernst & Young LLP（安永）', '2025年年报'],
        ['主要品牌', 'if（椰子水）、Innococo（电解质饮料）', 'IPO招股书'],
    ],
    'IFBH Limited 基本信息',
    [3.0, 7.0, 4.0]
)

add_fig('chart_ownership.png', 'IFBH 股权结构简图', 4.8)
doc.add_paragraph('1.2 历史沿革', style='Heading 2')
body('IFBH的创始人Pongsakorn Pongsak出身于泰国纺织巨头苏旺家族。2013年，他在泰国创立了"if"品牌，进入椰子水市场。品牌早期依托泰国本土的Nam Hom香水椰品种，以"100%纯椰子水"的产品定位快速打开市场。2015年，通过中国经销商进入大陆市场，借助电商渠道（天猫、京东）和新零售渠道（盒马、山姆）迅速扩张。2022年，公司将国际业务从General Beverage分拆，设立IFBH Pte. Ltd.作为上市主体。经过多轮融资后，于2025年6月在港交所成功上市。')

doc.add_paragraph('1.3 组织架构', style='Heading 2')
body('IFBH采用极致的轻资产组织架构。公司总部位于新加坡，在泰国设有运营团队，在中国大陆通过第三方经销商开展业务。董事会由8名董事组成：3名执行董事（含CEO）、1名非执行董事、4名独立非执行董事。董事会下设审核委员会（主席：Songvilai Jiraphothong女士）和薪酬及考核委员会（主席：Pathamakorn Buranasin女士），均由独立非执行董事担任主席。')

add_table(
    ['董事姓名', '职位', '类别', '委员会任职'],
    [
        ['Pongsakorn Pongsak', '行政总裁', '执行董事', '—'],
        ['Metaphon Pornanektana', '执行董事', '执行董事', '—'],
        ['Vipada Kanchanasorn', '执行董事', '执行董事', '—'],
        ['Tawat Kitkungvan', '非执行董事', '非执行董事', '—'],
        ['Thavee Thaveesangsakulthai', '独立非执行董事（主席）', '独立非执行', '审核/薪酬'],
        ['Songvilai Jiraphothong', '独立非执行董事', '独立非执行', '审核委员会主席/薪酬'],
        ['Pathamakorn Buranasin', '独立非执行董事', '独立非执行', '薪酬委员会主席/审核'],
        ['Supansa Kusonpattana', '独立非执行董事', '独立非执行', '—'],
    ],
    'IFBH董事会成员（截至2025年年报）',
    [3.5, 3.0, 2.5, 5.0]
)

doc.add_paragraph('1.4 行业地位', style='Heading 2')
body('根据西部证券等券商研报数据，IFBH在中国椰子水市场占有率连续五年位居第一。2024年市占率约34%（按零售额计算），主要竞品为Vita Coco（约18.5%）、Zico（约10.8%）及国内新兴品牌（合计约36.7%）。2025年受产品质量争议和竞品大规模入局影响，市占率从2024年Q4的62%骤降至2025年Q3的30.3%。全球范围内，IFBH在即饮椰子水市场排名第二（仅次于Vita Coco）。')

# ==================== 二、主营业务 ====================
pb()
doc.add_paragraph('二、主营业务分析', style='Heading 1')
ct.c(3)

doc.add_paragraph('2.1 商业模式', style='Heading 2')
body('IFBH采用极致的轻资产运营模式，核心特征是"三个100%外包"——100%原材料外包采购（通过General Beverage）、100%生产外包（12家泰国代工厂）、100%物流外包（第三方物流）。公司自身仅保留品牌管理、产品研发、品质监控和市场营销四项核心职能，由46名员工（其中销售及营销人员20名、仓储物流15名、研发5名、行政后台6名）完成全部运营。')

add_table(
    ['价值链环节', '负责方', '控制程度', '风险特征'],
    [
        ['椰子原料采购', 'General Beverage（控股股东）', '间接（通过GB）', '单一采集商依赖，从2位采集商+33农户收购'],
        ['产品灌装生产', '12家泰国代工厂', '间接（派驻监督团队）', '前5大代工厂占96.9%产量'],
        ['物流运输', '第三方物流商', '低（无直接合同关系）', '运输质量不可控'],
        ['中国分销', '2家核心经销商', '中（有年度框架协议）', '前5大客户贡献97.6%收入'],
        ['品牌管理', 'IFBH自身（20人销售团队）', '高（核心能力）', '品牌声誉极度依赖消费者信任'],
    ],
    'IFBH轻资产模式的价值链控制矩阵',
    [2.5, 3.5, 3.0, 5.0]
)

doc.add_paragraph('2.2 营收结构', style='Heading 2')
body('2025财年（截至2025年12月31日），公司实现营业收入1.76亿美元（同比增长11.9%），但净利润为2,277万美元（同比下降31.7%），呈现"增收不增利"态势。净利润下降主要由三方面因素驱动：一是上市相关专业费用约410万美元（一次性）；二是营销费用同比增长76%至1,302万美元；三是泰铢对美元升值侵蚀毛利率（从36.7%降至32.9%）。')

add_fig('chart_revenue.png', 'IFBH 营业收入与净利润趋势 (2023-2025)', 4.5)

add_table(
    ['指标', '2024财年', '2025财年', '同比变化', '信息来源'],
    [
        ['营业收入', '1.58亿美元', '1.76亿美元', '+11.9%', '2025年年报'],
        ['净利润（股东应占）', '3,332万美元', '2,277万美元', '-31.7%', '2025年年报'],
        ['毛利率', '36.7%', '32.9%', '-3.8pp', '2025年年报'],
        ['营销费用', '740万美元', '1,302万美元', '+76.0%', '2025年年报'],
        ['if品牌收入占比', '83.2%', '94.5%', '+11.3pp', '2025年年报'],
        ['中国内地收入占比', '92.4%', '90.4%', '-2.0pp', '2025年年报'],
        ['现金及等价物', '5,482万美元', '1.63亿美元', '+197%', '2025年年报（含IPO融资）'],
    ],
    'IFBH核心财务指标（2024-2025财年）',
    [3.0, 2.0, 2.5, 2.5, 4.0]
)

add_fig('chart_revenue_struct.png', 'IFBH 收入结构分布——按产品与按区域 (2025财年)', 5.0)
add_fig('chart_margin.png', 'IFBH 毛利率与营销费用率变化 (2024-2025)', 4.5)

doc.add_paragraph('2.3 核心资源与能力', style='Heading 2')
body('IFBH的核心竞争力在于三个方面。（1）品牌资产：if品牌在中国椰子水品类中具有最高的消费者认知度。（2）供应链关系：通过创始人家族控制的General Beverage，获得了泰国Nam Hom香水椰的稳定供应渠道——该品种树苗被泰国政府列为保护品种、禁止出口，构成了天然的原料壁垒。（3）渠道渗透：产品已进入天猫、京东等线上平台和山姆、盒马、Ole\'等线下高端零售渠道。')

body('然而，公司几乎不存在传统意义上的"核心技术"或"专利壁垒"。根据招股书和年报披露，公司未持有任何重大专利或专有技术，产品配方主要基于椰子水本身的天然属性，研发团队仅5人。公司的"护城河"本质上来源于品牌、渠道关系和原料供应链关系，而非技术创新。')

# ==================== 三、风险管理现状 ====================
pb()
doc.add_paragraph('三、风险管理现状', style='Heading 1')
ct.c(4)

doc.add_paragraph('3.1 现有风险管理体系', style='Heading 2')
body('根据IPO招股书和2025年年报的披露，IFBH目前的风险管理主要依托以下机制：（1）董事会层面的审核委员会负责监督风险管理框架的有效性；（2）管理层定期评估和报告主要风险因素；（3）在代工厂派驻小型品质监督团队；（4）通过第三方检测机构对产品进行定期抽检。')

body('然而，招股书同时坦承多项局限性：公司"对生产环节的直接监管有限"、"无法保证能与分销商续约"、"原料供应高度依赖控股股东General Beverage"。在IPO前，公司不存在正式的全面风险管理框架或首席风险官（CRO），风险管理工作主要由CEO和CFO直接负责，内控资源与46人的公司规模相匹配。上市后，公司承诺将逐步建立更完善的内控和风险管理体系。')

doc.add_paragraph('3.2 外部风险因素', style='Heading 2')
body('基于PESTEL分析框架，IFBH面临的主要外部风险因素如下：')

add_table(
    ['外部风险类别', '具体风险因素', '影响程度', '信息来源'],
    [
        ['政策法规风险', '中国食品安全标准更新（GB7718-2025/GB28050-2025）、进口食品注册新规（海关280号令）、广告用语限制', '高', '国家卫健委/海关总署/市监总局'],
        ['行业竞争风险', '椰树/菲诺/盒马等品牌大规模入局，价格带下沉至9.9元，IF市占率从62%骤降至30.3%', '高', '券商研报/电商平台数据'],
        ['社会文化风险', '消费者健康意识增强→对"100%纯椰子水"宣称的真实性更高要求；社交媒体负面传播速度加快', '高', '微博/小红书公开数据'],
        ['科技风险', '替代性功能性饮品（电解质水、植物基饮品）的技术进步可能分流消费者', '中', '行业分析报告'],
        ['自然环境风险', '泰国南部干旱/洪水灾害→Nam Hom香水椰减产；气候模式变化影响椰子产量和品质', '高', '泰国农业部门/世界银行'],
        ['宏观经济风险', '人民币对泰铢汇率波动→进口成本变化；中国消费信心下降→高端饮品需求减少', '中', '中国人民银行/外汇市场'],
    ],
    'IFBH外部风险因素分析（PESTEL框架）',
    [2.0, 5.0, 2.0, 5.0]
)

doc.add_paragraph('3.3 内部风险因素', style='Heading 2')

add_table(
    ['内部风险类别', '具体风险因素', '风险等级', '招股书/年报披露位置'],
    [
        ['战略风险', '过度依赖单一品类（椰子水占94.5%）和单一市场（中国占90.4%）', '重大', '招股书"风险因素"章节'],
        ['运营风险', '100%代工模式导致品控能力受限于第三方；经销商管理失控（副品牌因经销商问题停产数月）', '重大', '招股书/2025年报'],
        ['财务风险', '增收不增利（净利-31.7%）；应收账款激增173%；经营现金流暴跌70.5%', '重大', '2025年年报'],
        ['合规风险', '进口食品标签已受2次海关/市监局处罚；GB28050-2025"无添加"禁用语直接影响IF现有宣称', '高', '海关/市监局公开记录'],
        ['治理风险', '控股股东/供应商/客户三位一体（General Beverage）；仅46人团队的内控资源瓶颈', '高', '招股书/公司资料'],
    ],
    'IFBH内部风险因素分析',
    [2.0, 5.0, 2.0, 5.0]
)

# ==================== 四、风险识别与评估 ====================
pb()
doc.add_paragraph('四、风险识别与评估', style='Heading 1')
ct.c(5)

doc.add_paragraph('4.1 风险识别方法', style='Heading 2')
body('本报告采用"风险事件库法+财务指标分析法+舆情信号分析法"三种方法相结合的方式进行风险识别。风险事件库基于IFBH IPO招股书中"风险因素"章节（共披露32项风险因素）、2025年年报"管理层讨论与分析"章节中的风险提示，以及券商研报中的风险分析；财务指标分析法基于2025年年报中的核心财务比率（营收增速、毛利率变化、费用率、应收账款周转率等）识别财务反常信号；舆情信号分析法则通过公开的新闻、社交媒体和行业报告监测IF品牌的声誉风险和市场竞争力变化。')

doc.add_paragraph('4.2 风险矩阵评估', style='Heading 2')
body('使用风险矩阵（Risk Matrix）对各识别出的风险进行可能性和影响程度的二维评估。可能性分为五级（极低1-几乎不可能/低2-不太可能/中3-可能/高4-很可能/极高5-几乎确定），影响程度分为五级（极低1-可忽略/低2-轻微/中3-中等/高4-重大/极高5-灾难性）。综合评分=可能性×影响程度，≥15为重大风险、10-14为高风险、5-9为中风险、≤4为低风险。')

add_fig('chart_risk_matrix.png', 'IFBH 风险矩阵——8大风险的可能性与影响程度气泡图（气泡大小=综合评分）', 5.0)

add_table(
    ['编号', '风险名称', '可能性(1-5)', '影响(1-5)', '综合评分', '风险等级'],
    [
        ['R1', '供应链高度集中——原料100%依赖Nam Hom香水椰+100%通过General Beverage采购', '4（很可能）', '5（灾难性）', '20', '重大'],
        ['R2', '品牌声誉与食品安全——产品质量质疑+社交媒体负面声量55%+市占率暴跌', '4（很可能）', '5（灾难性）', '20', '重大'],
        ['R3', '单一市场/单一品类依赖——中国收入90.4%+椰子水产品94.5%', '5（几乎确定）', '4（重大）', '20', '重大'],
        ['R4', '财务恶化——增收不增利+费用失控+经营现金流恶化', '4（很可能）', '4（重大）', '16', '重大'],
        ['R5', '法律合规风险——标签法规升级+已受2次处罚+进口注册逾期', '4（很可能）', '3（中等）', '12', '高风险'],
        ['R6', '代工品控风险——100%外包+前5大代工厂占96.9%+监督能力有限', '3（可能）', '4（重大）', '12', '高风险'],
        ['R7', '汇率波动风险——泰铢成本/美元计价/人民币销售三重敞口', '4（很可能）', '2（轻微）', '8', '中风险'],
        ['R8', '关键人才依赖——CEO控制品牌+供应链两端+无正式继任计划', '2（不太可能）', '5（灾难性）', '10', '高风险'],
    ],
    'IFBH风险矩阵评估',
    [0.8, 5.0, 2.0, 2.0, 2.0, 2.0]
)

doc.add_paragraph('4.3 五大重大风险详细分析', style='Heading 2')

body('（一）R1 供应链高度集中风险。IFBH的供应链风险源于"三个100%"的极度集中：（1）原料端：100%依赖泰国Nam Hom香水椰品种，该品种仅产于叻丕府和龙仔厝府两个地区，树苗被泰国政府列为保护品种禁止出口。IFBH消耗量占泰国椰子水总产量的50%以上，供给端极为有限。（2）采购端：100%通过控股股东General Beverage采购椰子水原料——GB仅从2位采集商和33位农户处收集原料，任何一环中断都直接威胁IFBH的生产连续性。（3）生产端：100%外包给12家代工厂，前5大代工厂占96.9%的产量。泰国的干旱、洪水、病虫害或劳工短缺均可能同时影响原料供应和代工生产。IFBH自身无独立采购能力，缺乏对供应链的直接控制。')

body('（二）R2 品牌声誉与食品安全风险。2025年，多家媒体委托第三方检测机构对市售椰子水进行抽检，发现部分标称"100%椰子水"的产品存在外源水添加或糖浆添加的检测信号，IF产品位列其中。此次事件在微博等社交媒体迅速发酵，负面声量一度达到55%。天猫旗舰店评分从4.8降至4.5，差评关键词集中在"太甜""口感变淡""怀疑加糖"。与此同时，竞品Vita Coco等借机加大促销力度，IF市场份额从2024Q4的62%骤降至2025Q3的30.3%。')

body('（三）R3 单一市场/单一品类依赖风险。2025财年，中国内地市场贡献了90.4%的营业收入，if品牌椰子水产品贡献了94.5%的收入。一旦中国消费市场出现疲软（如经济下行周期），或椰子水品类被替代性饮品（如功能性饮料、电解质水）抢占市场份额，IFBH将面临无分散缓冲的极端风险。公司副品牌Innococo（电解质水）曾试图培育第二增长曲线，但因分销商管理问题停产数月，2025年收入暴跌63.2%，显示多元化战略执行能力不足。')

body('（四）R4 财务恶化风险。2025财年呈现典型的"增收不增利"反常模式——营收增长11.9%但净利润下降31.7%。核心驱动因素包括：营销费用激增76%至1,302万美元（远超营收增速），毛利率从36.7%降至32.9%（泰铢升值和原材料成本上升），以及IPO相关一次性费用410万美元。更具警示意义的是，公司经营现金流从2024年的正数转为负数，应收账款激增173%，暗示渠道压货或回款困难。尽管IPO融资后现金储备充裕（1.63亿美元），但若经营层面无法改善盈利质量，市值和融资能力将承压。')

body('（五）R5 法律合规风险。IF产品作为进口预包装食品，在中国市场须同时满足GB 7718（标签通则）、GB 28050（营养标签）和GB 2760（添加剂标准）等多项食品安全标准。2025年3月发布的新版GB 28050-2025和GB 7718-2025（2027年3月16日全面强制实施）对IF现有产品构成直接合规冲击：新规明确禁止使用"零添加""不添加""无添加糖"等宣传用语，而IF产品的核心卖点正是"100%纯椰子水""无添加糖"。此外，海关总署280号令（2026年6月1日已实施）要求境外食品生产企业在华注册，泰国代工厂的注册状态须立即确认。公司中国大陆总代理已因标签不合规被深圳海关罚款5.2万元、因许可证过期被深圳市监局罚款1万元。')

# ==================== 五、风险应对 ====================
pb()
doc.add_paragraph('五、风险应对策略', style='Heading 1')
ct.c(6)

doc.add_paragraph('5.1 应对策略总览', style='Heading 2')
body('针对上述五大重大风险，结合COSO ERM框架中的四类风险应对策略（规避/减轻/转移/接受），提出以下应对方案：')

add_table(
    ['风险编号', '风险名称', '主要应对策略', '具体措施', '预估成本/时间'],
    [
        ['R1', '供应链集中', '减轻+转移', '采购去依赖化（GB占比100%→70%→50%）；建立印尼/菲律宾椰子水测试采购线；签订长期锁价合同', '1,000-2,500万元/6-18个月'],
        ['R2', '品牌声誉', '减轻', '委托SGS等第三方全批次检测并公布结果；CEO公开回应；品牌溯源营销；投保产品召回险', '11万元（检测3万+保险8万）/立即'],
        ['R3', '单一市场依赖', '减轻', '拓展东南亚/澳洲等海外市场；培育第二增长曲线；维持中国市场基本盘', '营销预算内/持续'],
        ['R4', '财务恶化', '减轻+接受', '建立200-300万元风险准备金；优化营销费用效率；监控经营现金流/应收款', '200-300万元准备金/1个月'],
        ['R5', '法律合规', '规避+减轻', '立即确认泰国工厂在华注册编号；全面审查标签合规性；停止使用"无添加糖"宣传语；聘请食品合规法律顾问', '5-10万元/2周-3个月'],
    ],
    '五大重大风险的应对策略矩阵',
    [1.0, 2.0, 2.5, 5.0, 3.5]
)

doc.add_paragraph('5.2 优先行动时间线', style='Heading 2')
body('基于风险的紧迫度和影响程度，建议按以下时间线推进应对措施：')
body('紧急响应（0-2周）：确认泰国工厂在华注册状态（海关280号令已逾期）；全面审查现有产品标签合规性，立即停用"无添加糖""零添加"等被禁用语；委托第三方检测机构出具全批次检测报告并公开发布。')
body('短期整改（1-3个月）：向海关总署完成境外生产企业注册补办；聘请食品合规法律顾问；建立标签预审核和证照到期自动提醒SOP；启动与3家新泰国采集商的合作协议谈判。')
body('中期巩固（3-12个月）：建立印尼/菲律宾椰子水小规模测试采购线（100-300万元）；代工厂产能分散化（300-800万元）；建立200万元风险准备金；投保产品召回险（500万元保额）。')

doc.add_paragraph('5.3 对现有策略的评价与改进建议', style='Heading 2')
body('IFBH在招股书中已承认上述风险并提出了部分缓释措施（如降低对General Beverage的采购依赖至70%→50%），但执行进度和效果待观察。相较于同体量的上市消费品公司，IFBH的风险应对存在三个主要不足：（1）过度依赖家族关系而非制度化机制来管理供应链风险；（2）缺乏正式的危机管理预案和模拟演练；（3）未建立独立的风险管理职能部门。建议公司在2026年内设立首席风险官（CRO）或风险管理委员会，将风险管理从CEO个人职责升级为机构化职能。')

# ==================== 六、内控基础 ====================
pb()
doc.add_paragraph('六、内部控制基础', style='Heading 1')
ct.c(7)

doc.add_paragraph('6.1 治理结构', style='Heading 2')
body('IFBH的治理结构遵循港交所上市规则的要求，设立了审核委员会和薪酬及考核委员会，均由独立非执行董事担任主席。8名董事会成员中有4名独立非执行董事（占比50%），形式上满足港交所《企业管治守则》对独立性的基本要求。但需注意两个治理风险点：（1）创始人家族通过General Beverage持有65.51%的绝对控股权，中小股东在公司重大决策中的话语权有限；（2）General Beverage同时是IFBH的控股股东、核心供应商和客户，关联交易的定价公允性和审批独立性存在潜在冲突。')

doc.add_paragraph('6.2 主要内控制度', style='Heading 2')
body('根据招股书和年报的有限披露，IFBH目前的内控制度主要包括：（1）代工厂品质监督——向代工厂派驻小型团队监督生产标准执行；（2）第三方产品检测——定期委托检测机构对产品进行抽检；（3）关联交易审批——General Beverage与IFBH之间的关联交易须经独立非执行董事审核。然而，对于一家年营收超1.7亿美元的上市消费品公司而言，仅46人的团队规模意味着内控资源严重不足。招股书坦承"对生产环节的直接监管有限"，且2025年年报显示公司仍在"建立和完善内控体系"的过程中。')

doc.add_paragraph('6.3 内控缺陷与改进建议', style='Heading 2')
body('基于公开信息的分析，识别出以下内控改进空间：（1）设立独立的内审部门——目前内审职能由审核委员会直接承担，但缺乏专职内审人员；（2）建立供应商和经销商管理制度——对前5大供应商/客户的依赖度极高，须建立独立的信用评估和定期审计机制；（3）完善信息沟通机制——46人规模的扁平化组织在成长期效率高，但随着业务复杂度的提升，正式的信息传递和审批流程需要建立或强化；（4）提升关联交易透明度——General Beverage的多重角色（控股股东/供应商/客户），建议在年报和中期报告中更详细地披露关联交易的定价依据和独立董事审批意见。')

# ==================== 七、监督与改进 ====================
pb()
doc.add_paragraph('七、监督与改进', style='Heading 1')
ct.c(8)

doc.add_paragraph('7.1 监督主体', style='Heading 2')
body('IFBH的监督体系由四个层面构成：（1）董事会审核委员会——负责监督风险管理框架和内控体系的有效性，由3名独立非执行董事组成；（2）外部审计师——Ernst & Young LLP（安永）作为独立审计师，对年度财务报表进行审计并出具审计意见；（3）港交所及香港证监会——作为上市公司，IFBH须遵守《上市规则》和《证券及期货条例》的持续披露义务；（4）媒体与公众监督——2025年多家媒体的产品检测报道直接触发了品牌危机，显示了外部监督的力量。')

doc.add_paragraph('7.2 信息披露机制', style='Heading 2')
body('IFBH的信息披露主要通过以下渠道：港交所披露易平台（年报、中期报告、关联交易公告、股东周年大会通告等）、公司官网及投资者关系页面。2025年年报中关于风险因素的披露较为详细（涵盖32项风险因素），但在以下方面可进一步改进：（1）关联交易的定价细节披露不足——General Beverage作为核心供应商的采购金额和定价基准未充分量化；（2）ESG信息披露空白——公司未发布独立的ESG报告，环境（如椰子种植的水资源消耗）、社会（如供应链劳工权益）和治理（如反腐败政策）维度的信息缺失；（3）风险应对措施的执行进度缺乏定量跟踪披露。')

doc.add_paragraph('7.3 持续改进机制', style='Heading 2')
body('基于上述分析，建议IFBH建立以下持续改进机制：（1）每年更新企业风险登记册（Risk Register），由审核委员会审批并向董事会报告；（2）引入ISO 31000风险管理标准作为内部风险管理框架的基础，逐步从"被动披露风险"转型为"主动管理风险"；（3）在2026年年报中加入风险应对措施的定量进度披露（如：GB采购占比从100%降至X%，新采集商数量从2增至X，产品检测合格率从X%升至X%）；（4）聘请独立第三方进行年度内控有效性评估并公开发布评估结论。')

# ==================== 八、总结与展望 ====================
pb()
doc.add_paragraph('八、总结与展望', style='Heading 1')
ct.c(9)

doc.add_paragraph('8.1 风险管理总体评价', style='Heading 2')
body('IFBH作为一家新上市的快消品企业，在风险管理和内控体系建设方面处于"起步阶段"。公司的优势在于：创始人深度参与、组织扁平决策高效、品牌在细分赛道的市场地位稳固、IPO融资后现金储备充足。公司的劣势在于：供应链和市场的双重集中度极高（原料100%+市场90%+品类95%）、内控资源与业务规模不匹配（46人管理1.76亿美元营收）、关联交易治理存在结构性利益冲突。')

doc.add_paragraph('8.2 主要风险敞口与前瞻展望', style='Heading 2')
body('展望2026-2027年，IFBH面临三个关键挑战窗口期：（1）2026年下半年——泰国雨季对椰子产量的实际影响将决定原材料成本走向，公司须在此窗口期内实质性推进采购多元化（GB占比降至70%以下）；（2）2027年3月——GB 7718-2025和GB 28050-2025全面强制实施，IF须在18个月过渡期内完成标签全面换版，否则将面临大规模退货或处罚风险；（3）2025-2027年——竞品大规模入局和价格战将持续考验IF的品牌溢价能力和市场份额。未来12-18个月，将是检验IFBH风险管理能力的关键窗口期。')

doc.add_paragraph('8.3 改进空间', style='Heading 2')
body('（1）制度层面——设立首席风险官或风险管理委员会，将风险管理从CEO个人职责升级为机构化职能。（2）流程层面——建立季度风险自评（RCSA）流程，6大风险维度由对应部门负责人定期自评并上报。（3）技术层面——引入数字化风险管理工具（如本项目开发的Risk Sentinel智能体），实现风险信号的自动化监测和预警。（4）文化层面——在46人团队中建设"人人都是风险管理者"的组织文化，将风险意识嵌入日常运营决策。')

# ==================== 附录 ====================
pb()
doc.add_paragraph('附  录', style='Heading 1')
ct.c(10)

doc.add_paragraph('附录A  关键财务数据表', style='Heading 2')
add_table(
    ['指标', '2023财年', '2024财年', '2025财年', '同比变化(2025 vs 2024)'],
    [
        ['营业收入（百万美元）', '约118', '158', '176', '+11.9%'],
        ['净利润（百万美元）', '约25', '33.3', '22.8', '-31.7%'],
        ['毛利率', '约38%', '36.7%', '32.9%', '-3.8pp'],
        ['椰子水收入占比', '约90%', '83.2%', '94.5%', '+11.3pp'],
        ['中国内地收入占比', '约95%', '92.4%', '90.4%', '-2.0pp'],
        ['员工人数', '—', '46', '46', '—'],
        ['营销费用率', '—', '约5%', '7.4%', '+2.4pp'],
    ],
    'IFBH核心财务数据（2023-2025财年）',
    [3.0, 2.0, 2.5, 2.5, 4.0]
)

doc.add_paragraph('附录B  主要信息来源清单', style='Heading 2')

body_ni('1. IFBH Limited. "聆讯后资料集"（IPO招股书）. 香港联交所披露易, 2025年6月.  https://www1.hkexnews.hk/listedco/listconews/sehk/2025/0620/2025062000034_c.pdf')
body_ni('2. IFBH Limited. "截至2025年12月31日止财政年度的年度业绩公告". 香港联交所披露易, 2026年2月26日.')
body_ni('3. IFBH Limited. "2025年中期报告". 香港联交所披露易, 2025年9月.')
body_ni('4. IFBH Limited. "股东周年大会通告". 香港联交所披露易, 2026年3月30日.')
body_ni('5. 西部证券. "IFBH(6603.HK) 椰风乘势千帆竞，龙头领航拓新机"（研报）. 2025年9月.')
body_ni('6. 新京报. "46人撑起IFBH上市神话：九成收益靠椰子水 代工模式藏风险". 2025年.')
body_ni("7. 每日经济新闻. \"中国市场'喝'出一个椰子水港股IPO——供应链地域分布集中、单品类规模受限是if椰子水未来考验\". 2025年4月.")
body_ni('8. 证券时报. "if椰子水母公司IFBH拟港股IPO 供应链延展待观察". 2025年.')
body_ni('9. COSO. "Enterprise Risk Management — Integrating with Strategy and Performance". 2017.')
body_ni('10. ISO. "ISO 31000:2018 Risk management — Guidelines". 2018.')
body_ni('11. 国家卫生健康委员会. "GB 7718-2025 预包装食品标签通则". 2025年3月发布.')
body_ni('12. 国家卫生健康委员会. "GB 28050-2025 预包装食品营养标签通则". 2025年3月发布.')
body_ni('13. 海关总署. "第280号令：进口食品境外生产企业注册管理规定". 2025年10月发布, 2026年6月1日实施.')
body_ni('14. 市场监管总局. "第100号令：食品标识监督管理办法". 2025年3月发布.')
body_ni('15. 联交所. "香港联合交易所有限公司证券上市规则——附录C1：企业管治守则".')

# ===== 结尾 =====
body_ni('')
p=doc.add_paragraph('— 报告完 —', style='Normal'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER

# ===== 保存 =====
out = r'C:\Users\23994\risk_sentinel\IFBH_Enterprise_Risk_Report_v2.docx'
doc.save(out)
print(f'Saved: {out}')
print(f'Size: {os.path.getsize(out)} bytes')
